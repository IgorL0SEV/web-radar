#!/usr/bin/env python
"""
web-radar parse_docs — обработка .docx и .xlsx файлов.

Команды:
  python scripts/parse_docs.py docx <path>  — извлечь текст из .docx
  python scripts/parse_docs.py xlsx <path> [--sheet <name|index>] [--max-rows <N>]  — извлечь таблицу из .xlsx

Результат выводится в stdout как JSON.
"""

import argparse
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]


def parse_docx(path: str):
    """Извлечь текст и таблицы из .docx файла."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "Модуль python-docx не установлен. Установите: pip install python-docx"}

    file_path = Path(path)
    if not file_path.is_absolute():
        file_path = ROOT / file_path

    if not file_path.exists():
        return {"error": f"Файл не найден: {file_path}"}

    try:
        doc = Document(str(file_path))
    except Exception as e:
        return {"error": f"Ошибка чтения .docx: {e}"}

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    return {
        "source": str(file_path),
        "type": "docx",
        "paragraphs": paragraphs,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "tables": tables,
    }


def parse_xlsx(path: str, sheet=None, max_rows=100):
    """Извлечь данные из .xlsx файла."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"error": "Модуль openpyxl не установлен. Установите: pip install openpyxl"}

    file_path = Path(path)
    if not file_path.is_absolute():
        file_path = ROOT / file_path

    if not file_path.exists():
        return {"error": f"Файл не найден: {file_path}"}

    try:
        wb = load_workbook(str(file_path), read_only=True, data_only=True)
    except Exception as e:
        return {"error": f"Ошибка чтения .xlsx: {e}"}

    sheet_names = wb.sheetnames
    target_sheet = None

    if sheet is not None:
        if isinstance(sheet, int):
            if 0 <= sheet < len(sheet_names):
                target_sheet = wb[sheet_names[sheet]]
            else:
                return {"error": f"Лист с индексом {sheet} не найден. Доступные: {sheet_names}"}
        else:
            if sheet in sheet_names:
                target_sheet = wb[sheet]
            else:
                return {"error": f"Лист '{sheet}' не найден. Доступные: {sheet_names}"}
    else:
        target_sheet = wb[sheet_names[0]]

    rows = []
    for i, row in enumerate(target_sheet.iter_rows(values_only=True)):
        if i >= max_rows:
            break
        rows.append([str(cell) if cell is not None else "" for cell in row])

    wb.close()

    return {
        "source": str(file_path),
        "type": "xlsx",
        "sheet_names": sheet_names,
        "active_sheet": target_sheet.title,
        "row_count": len(rows),
        "max_rows_limit": max_rows,
        "data": rows,
    }


def main():
    parser = argparse.ArgumentParser(description="web-radar parse_docs — обработка документов")
    sub = parser.add_subparsers(dest="command")

    p_docx = sub.add_parser("docx", help="Извлечь текст из .docx")
    p_docx.add_argument("path", help="Путь к .docx файлу")

    p_xlsx = sub.add_parser("xlsx", help="Извлечь таблицу из .xlsx")
    p_xlsx.add_argument("path", help="Путь к .xlsx файлу")
    p_xlsx.add_argument("--sheet", help="Имя или индекс листа (по умолчанию первый)")
    p_xlsx.add_argument("--max-rows", type=int, default=100, help="Максимальное количество строк")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    if args.command == "docx":
        result = parse_docx(args.path)
    elif args.command == "xlsx":
        sheet = args.sheet
        if sheet is not None and sheet.isdigit():
            sheet = int(sheet)
        result = parse_xlsx(args.path, sheet=sheet, max_rows=args.max_rows)
    else:
        result = {"error": f"Неизвестная команда: {args.command}"}

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()